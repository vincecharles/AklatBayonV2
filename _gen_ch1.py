"""Generate Chapter 1 - Introduction for AklatBayon capstone."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(2.54)

# Style setup
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14) if level == 0 else Pt(12)
    p.paragraph_format.space_before = Pt(24) if level > 0 else Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 2.0
    return p

def paragraph(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    return p

def subheading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 2.0
    return p

# ── Chapter Title ──────────────────────────────────────
heading('CHAPTER 1', level=0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('THE PROBLEM AND ITS BACKGROUND')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

# ── 1.1 Background of the Study ───────────────────────
heading('1.1 Background of the Study')

paragraph(
    'Libraries in Philippine universities have been trying to keep up with the shift toward digital systems for years now. '
    'Some schools already made the jump, but a lot of them, especially the smaller or older institutions, are still relying on '
    'manual processes or semi-automated setups that barely get the job done. At FEATI University, the library has been dealing '
    'with the same kind of growing pains. Students line up to borrow books, staff manually write down transactions on logbooks, '
    'and tracking overdue materials is more of a guessing game than an actual system.'
)

paragraph(
    'The thing is, it does not have to be this way. Library management systems have been around for a while, and there are '
    'commercial options like Koha, Destiny, and LibSys that bigger institutions use. But the problem with most of these is '
    'cost, complexity, or both. They require dedicated servers, trained IT personnel for maintenance, and licensing fees that '
    'can pile up over time. For a university like FEATI, those are real barriers.'
)

paragraph(
    'That is where this study comes in. The researchers saw an opportunity to build something that fits the specific needs '
    'of the FEATI University Library without the overhead of enterprise solutions. The idea was to develop a web-based Library '
    'Management System called AklatBayon, a name that combines the Filipino words "Aklat" (book) and "Bayon" (to carry or '
    'bring along), reflecting the goal of making library resources more accessible and portable for the FEATI community.'
)

paragraph(
    'AklatBayon was designed from the ground up as a modern, cloud-hosted web application deployed on Netlify with a '
    'Neon PostgreSQL database backend. Unlike traditional library systems that need physical servers, this setup runs entirely '
    'on serverless infrastructure, which means zero server maintenance on the university\'s end. The system covers the core '
    'operations of a university library: cataloging books using the Library of Congress Classification system, managing '
    'circulation (borrowing, returning, and renewing), handling fines automatically, processing book reservations, tracking '
    'attendance through RFID, and generating reports.'
)

paragraph(
    'What makes AklatBayon different from a generic library tool is that it was built around the actual organizational structure '
    'of FEATI. The system recognizes seven distinct user roles: System Administrator, Head Librarian, Librarian Staff, Faculty '
    '(with subtypes for Teaching, Non-Teaching, and Department Chair), Student, Student Assistant, and Guest. Each of these roles '
    'has a different set of permissions and borrowing policies. For example, a Head Librarian can borrow up to 15 books for 60 days, '
    'while a student is limited to 3 books for 7 days. These rules were not just coded in arbitrarily. They reflect the actual policies '
    'that the FEATI Library follows.'
)

paragraph(
    'The system also integrates with the Library of Congress online database, allowing librarians to search and import book records '
    'directly instead of typing everything from scratch. This is a practical feature that saves time and reduces data entry errors, '
    'especially for cataloging newly acquired books.'
)

# ── 1.2 Statement of the Problem ──────────────────────
heading('1.2 Statement of the Problem')

paragraph(
    'The FEATI University Library currently operates without a unified digital system for managing its day-to-day functions. '
    'Book transactions are tracked manually or through disconnected spreadsheets, which creates several problems that affect '
    'both the library staff and the students who depend on the library\'s services.'
)

paragraph(
    'Based on the observations and preliminary assessment conducted by the researchers, the following specific problems were identified:'
)

# Problem list
problems = [
    'How can the library reduce the time and errors involved in manual book borrowing, returning, and renewal processes?',
    'How can the system enforce role-based borrowing policies automatically, considering that FEATI has seven different user roles with varying book limits, loan durations, and fine rates?',
    'How can overdue books and unpaid fines be tracked and computed without requiring library staff to check records one by one?',
    'How can the library provide an accessible Online Public Access Catalog (OPAC) that allows students and faculty to search the collection remotely?',
    'How can book reservations be managed fairly through a queue system, ensuring that materials with high demand are distributed equitably among borrowers?',
    'How can library attendance be monitored efficiently using RFID technology instead of manual sign-in sheets?',
    'How can the system maintain a reliable audit trail of all transactions, user actions, and system events for accountability and reporting purposes?',
]

for i, prob in enumerate(problems, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'{i}. {prob}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ── 1.3 Objectives of the Study ───────────────────────
heading('1.3 Objectives of the Study')

subheading('General Objective')

paragraph(
    'The main goal of this study is to develop and deploy a web-based Library Management System for FEATI University '
    'that digitizes and automates the core library operations, making them more efficient, accurate, and accessible '
    'to all members of the university community.'
)

subheading('Specific Objectives')

paragraph('Specifically, this study aims to:', indent=False)

objectives = [
    'Design and implement a role-based access control system with seven user roles and 23 granular permissions '
    'organized across eight permission groups (Users, Students, Catalog, Circulation, Finance, Reports, System, and General) '
    'to ensure that each user can only access the features appropriate to their position.',

    'Develop a book cataloging module that supports the Library of Congress Classification (LCC) system with 21 main classes '
    'and integrates with the Library of Congress online API for importing book records.',

    'Build a circulation module that handles book issuing, returning, and renewal with role-based borrowing policies, including '
    'configurable loan durations, maximum book limits, renewal caps, and fine rates for each of the seven borrower categories.',

    'Implement an automated fine computation system that detects overdue books and calculates penalties based on the borrower\'s role, '
    'with a configurable fine block threshold that prevents borrowing when unpaid fines exceed a set amount.',

    'Create a book reservation system with a first-in, first-out (FIFO) queue, a 48-hour claim window, and automatic expiration '
    'for unclaimed reservations.',

    'Integrate RFID-based attendance tracking for monitoring library foot traffic, with support for real-time tap logging '
    'and role-based attendance reports.',

    'Deploy the application on a serverless cloud infrastructure using Netlify for hosting, Neon PostgreSQL for the database, '
    'and Drizzle ORM for type-safe database operations, ensuring the system requires no physical server maintenance.',

    'Develop a comprehensive audit logging system that records all user actions, including login and logout events, data '
    'modifications, and system configuration changes for transparency and accountability.',
]

for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'{i}. {obj}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ── 1.4 Significance of the Study ─────────────────────
heading('1.4 Significance of the Study')

paragraph(
    'This study holds practical value for several groups within and beyond FEATI University. '
    'The researchers believe that the development and deployment of AklatBayon will benefit the following stakeholders:'
)

subheading('To the FEATI University Library Staff')
paragraph(
    'The system directly addresses the daily pain points of the library personnel. Instead of manually writing '
    'transaction records in logbooks, the staff can process book issuances and returns in a few clicks. The automated '
    'fine computation removes the burden of calculating overdue penalties by hand, and the reservation queue system '
    'eliminates the need to manually track who requested which book first. For the Head Librarian, the reports and '
    'inventory modules provide a clear picture of the library\'s status at any given time without having to count books physically.'
)

subheading('To the Students and Faculty')
paragraph(
    'Students and faculty gain access to an Online Public Access Catalog that lets them search the library\'s collection '
    'from anywhere with an internet connection. They no longer need to physically visit the library just to check if a book '
    'is available. The borrower profile with QR code integration also makes identification faster during transactions. '
    'Faculty members, in particular, benefit from the book recommendation feature, which allows them to suggest titles '
    'for acquisition without going through a paper-based request process.'
)

subheading('To the University Administration')
paragraph(
    'For the administration, AklatBayon provides a layer of accountability through its audit logging system. Every login, '
    'every book transaction, every configuration change is recorded with a timestamp and the identity of the user who performed '
    'the action. This kind of transparency is important for institutional governance. The cloud-based deployment also means '
    'the university does not need to invest in dedicated servers or hire additional IT staff to maintain the system.'
)

subheading('To Future Researchers')
paragraph(
    'This study serves as a reference for future capstone or thesis projects that involve library automation, '
    'role-based access control, or serverless web application development. The documented architecture, covering '
    'the frontend design, serverless API layer, and PostgreSQL database schema with 16 tables and 11 foreign key '
    'relationships, can be used as a template or starting point for similar projects in other educational institutions.'
)

# ── 1.5 Scope and Limitations ─────────────────────────
heading('1.5 Scope and Limitations')

subheading('Scope')

paragraph(
    'The AklatBayon Library Management System covers the following functional areas:'
)

scope_items = [
    'User management with seven roles (System Administrator, Head Librarian, Librarian Staff, Faculty, Student, '
    'Student Assistant, and Guest) and 23 permissions distributed across eight groups.',

    'Book catalog management with support for authors, publishers, hierarchical categories, ISBN tracking, '
    'call number assignment based on the Library of Congress Classification, and barcode label generation using the CODE128 format.',

    'Integration with the Library of Congress online API for searching and importing book records into the local catalog.',

    'A circulation module covering book issuance, return processing, and renewals with role-specific policies: '
    'Head Librarian (15 books, 60 days, 3 renewals), Librarian Staff (10 books, 30 days, 2 renewals), '
    'Faculty Teaching (10 books, 30 days, 2 renewals), Faculty Non-Teaching (5 books, 14 days, 1 renewal), '
    'Department Chair (10 books, 30 days, 2 renewals), Student (3 books, 7 days, 1 renewal), and '
    'Student Assistant (3 books, 7 days, 1 renewal).',

    'Automated overdue fine computation with configurable daily rates per role and a fine block threshold of 100 pesos '
    'that prevents further borrowing until fines are settled.',

    'A book reservation system with a FIFO queue, a 48-hour claim window, automatic expiration, and reservation '
    'limits (2 for students, 5 for faculty).',

    'RFID-based attendance tracking with tap-time logging.',

    'An audit log that records all system actions including user logins, logouts, data creation, modification, and deletion.',

    'A configurable settings module with over 30 parameters covering borrowing policies, fine rates, and reservation rules.',

    'Cloud deployment on Netlify with serverless functions (Netlify Functions) and a Neon PostgreSQL database accessed through Drizzle ORM.',

    'The system is accessible through any modern web browser and consists of 23 application pages covering all the modules listed above.',
]

for i, item in enumerate(scope_items, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'{i}. {item}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

subheading('Limitations')

paragraph(
    'While the system covers the essential functions of a university library, there are areas that fall outside '
    'the current scope of this study:'
)

limitations = [
    'The system does not implement advanced security features such as password hashing, JWT-based session tokens, '
    'server-side authentication middleware, or rate limiting. These are planned for a future phase but are not included '
    'in the current version.',

    'The RFID attendance feature relies on manual input simulation. Integration with physical RFID hardware (readers and tags) '
    'was not tested due to the unavailability of equipment during development.',

    'The system does not support traditional library interoperability protocols such as MARC records, Z39.50, or SRU/SRW. '
    'Catalog data exchange is currently limited to the Library of Congress REST API.',

    'Multi-branch library support is not included. The system is designed for a single-library, single-campus setup.',

    'Email or SMS notifications for reservation availability, overdue reminders, and fine alerts are not yet implemented. '
    'Notifications currently appear only within the web application.',

    'The system has not undergone formal usability testing with actual library staff and students at FEATI University. '
    'Feedback collection and user acceptance testing are scheduled as part of the next development phase.',
]

for i, item in enumerate(limitations, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'{i}. {item}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ── 1.6 Definition of Terms ───────────────────────────
heading('1.6 Definition of Terms')

paragraph(
    'The following terms are used throughout this study. Definitions are provided for clarity, '
    'particularly for readers who may not be familiar with library science or web development terminology.'
)

terms = [
    ('AklatBayon',
     'The name of the Library Management System developed in this study. It is a combination of the Filipino words '
     '"Aklat" (book) and "Bayon" (to carry). It refers to the web application deployed at aklatbayonfeati.netlify.app.'),

    ('Circulation',
     'The process of lending library materials to borrowers and receiving them back. In the context of this system, '
     'circulation includes book issuance, return processing, and renewal transactions.'),

    ('Drizzle ORM',
     'A lightweight, type-safe Object-Relational Mapping tool for TypeScript that the system uses to interact with '
     'the PostgreSQL database. It generates SQL queries from TypeScript code and handles schema migrations.'),

    ('FIFO (First-In, First-Out)',
     'A queue processing method where the first person to make a reservation is the first to be notified when the '
     'book becomes available. The system uses FIFO for its reservation queue.'),

    ('Library of Congress Classification (LCC)',
     'A classification system used by most academic and research libraries in the United States and internationally. '
     'It organizes books by subject using a combination of letters and numbers. The system implements all 21 main LCC classes.'),

    ('Neon PostgreSQL',
     'A serverless PostgreSQL database service. The system uses Neon as its cloud database, accessed through Netlify\'s '
     'built-in Neon integration. It stores all 16 tables of the application\'s data.'),

    ('Netlify',
     'A cloud platform for deploying web applications. The system is hosted on Netlify, which provides static site hosting '
     'and serverless function execution without the need for a traditional web server.'),

    ('Netlify Functions',
     'Serverless functions that run on Netlify\'s infrastructure. The system uses two Netlify Functions: one for handling '
     'all REST API requests (api.mts) and another for seeding the database with initial data (seed.mts).'),

    ('OPAC (Online Public Access Catalog)',
     'A publicly accessible digital catalog that allows users to search for books by title, author, ISBN, or subject. '
     'The system\'s OPAC is accessible through the catalog page even without a full user login.'),

    ('RBAC (Role-Based Access Control)',
     'A security model where access to system features is determined by the user\'s assigned role rather than by '
     'individual user settings. The system uses RBAC with seven roles and 23 permissions.'),

    ('RFID (Radio-Frequency Identification)',
     'A technology that uses electromagnetic fields to identify and track tags attached to objects or cards. '
     'In the system, RFID is used for user identification during attendance logging and as an alternative login method.'),

    ('Serverless Architecture',
     'A cloud computing model where the cloud provider manages the server infrastructure. The application code runs '
     'in stateless functions that are triggered by HTTP requests. The system uses this model through Netlify Functions '
     'and Neon PostgreSQL, eliminating the need for a dedicated application server.'),

    ('SweetAlert2',
     'A JavaScript library used in the system for displaying styled dialog boxes, confirmation prompts, and '
     'notification messages instead of the browser\'s default alert boxes.'),
]

for term, definition in terms:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.27)
    run_term = p.add_run(f'{term} ')
    run_term.bold = True
    run_term.font.name = 'Times New Roman'
    run_term.font.size = Pt(12)
    run_def = p.add_run(f'- {definition}')
    run_def.font.name = 'Times New Roman'
    run_def.font.size = Pt(12)

# Save
output = r'c:\Users\Bins\AklatBayonV2\Chapter1_AklatBayon.docx'
doc.save(output)
print(f'Saved: {output}')
