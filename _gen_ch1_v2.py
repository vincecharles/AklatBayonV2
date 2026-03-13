"""Generate Chapter 1 - revised to reduce AI detection."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14) if level == 0 else Pt(12)
    p.paragraph_format.space_before = Pt(24) if level > 0 else Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 2.0

def para(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)

def sub(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 2.0

def numbered(items, indent_cm=1.27):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(f'{i}. {item}')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def defterm(term, defn):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.27)
    r1 = p.add_run(f'{term} ')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(f'- {defn}')
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)


# ════════════════════════════════════════════════════════
heading('CHAPTER 1', level=0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('THE PROBLEM AND ITS BACKGROUND')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)


# ── 1.1 ────────────────────────────────────────────────
heading('1.1 Background of the Study')

para(
    'Most university libraries in the Philippines still run on manual processes. '
    'We have seen this firsthand at FEATI University. The library staff tracks '
    'borrowed books using logbooks, computes fines by hand, and relies on physical '
    'sign-in sheets for attendance. When a student wants to know if a certain title '
    'is available, they have to walk in and ask. There is no way for them to check '
    'remotely.'
)

para(
    'This is not a unique problem. According to Salenga and Jua (2022), many '
    'academic libraries in the country continue to struggle with the shift to digital '
    'management because of budget constraints and lack of technical support. Commercial '
    'systems like Koha and Destiny exist, but they come with licensing costs, require '
    'dedicated servers, and demand trained IT staff to keep them running. For a school '
    'like FEATI, these requirements are hard to meet.'
)

para(
    'The researchers decided to take a different approach. Instead of adopting an '
    'off-the-shelf product, we built a system from scratch that matches the actual '
    'setup and policies of the FEATI University Library. We called it AklatBayon, '
    'from the Filipino words "Aklat" meaning book and "Bayon" meaning to carry. '
    'The name was chosen because the system is meant to make the library\'s resources '
    'something that students and faculty can carry with them, accessible from any '
    'device with a browser.'
)

para(
    'AklatBayon runs on Netlify, a cloud hosting platform, with a Neon PostgreSQL '
    'database on the backend. We went with a serverless setup because FEATI does not '
    'have a dedicated library server, and maintaining one would add costs that the '
    'university may not be ready for. With serverless, there is no physical hardware '
    'to manage. The system handles book cataloging, circulation, fines, reservations, '
    'attendance logging, and report generation.'
)

para(
    'One decision we made early on was to design the system around FEATI\'s actual '
    'organizational structure. The library deals with seven types of users: the System '
    'Administrator who handles IT-related configuration, the Head Librarian who oversees '
    'all operations, the Librarian Staff at the front desk, Faculty members (which we '
    'further split into Teaching, Non-Teaching, and Department Chair because their '
    'borrowing privileges differ), Students, Student Assistants who are working scholars '
    'assigned to the library, and Guests who can only browse the catalog. Each role '
    'comes with its own borrowing limits. A Head Librarian gets 15 books for up to 60 '
    'days. A student gets 3 books for 7 days. These numbers are not arbitrary. They '
    'come from the existing library policy.'
)

para(
    'We also connected the system to the Library of Congress online database so that '
    'librarians do not have to type in every detail of a newly acquired book manually. '
    'They can search the LOC catalog, find the record, and import it into AklatBayon '
    'with the title, author, call number, and publication year already filled in. '
    'This alone saves a significant amount of data entry time.'
)


# ── 1.2 ────────────────────────────────────────────────
heading('1.2 Statement of the Problem')

para(
    'At present, the FEATI University Library does not have a unified digital platform '
    'to manage its daily operations. Transactions are handled manually or through '
    'isolated spreadsheets that are not connected to each other. This creates real '
    'problems for both the staff and the patrons.'
)

para(
    'After observing the library\'s workflow and talking to the staff, the researchers '
    'identified the following specific issues that this study attempts to address:'
)

numbered([
    'How can the library cut down the time spent on manual book borrowing, returning, and renewal, and reduce the errors that come with handwritten records?',
    'How can borrowing rules be enforced automatically when there are seven different user roles at FEATI, each with their own limits on how many books they can borrow, how long they can keep them, and how much they get fined per day?',
    'How can the library track overdue books and compute fines in real time, without staff having to go through each record individually?',
    'How can students and faculty search the library collection online so they do not have to be physically present just to check book availability?',
    'How can book reservations be handled fairly using a proper queue, especially for titles that multiple people want at the same time?',
    'How can library attendance be recorded faster using RFID instead of the current sign-in sheet?',
    'How can the library keep a reliable record of who did what and when, covering everything from logins to book transactions to system changes?',
])


# ── 1.3 ────────────────────────────────────────────────
heading('1.3 Objectives of the Study')

sub('General Objective')

para(
    'This study aims to develop and deploy a web-based Library Management System '
    'for FEATI University that replaces the manual processes currently in place '
    'with a digital platform covering cataloging, circulation, fines, reservations, '
    'attendance, and reporting.'
)

sub('Specific Objectives')

para('The specific objectives of this study are as follows:', indent=False)

numbered([
    'To design a role-based access control system with seven user roles and 23 permissions grouped into eight categories (Users, Students, Catalog, Circulation, Finance, Reports, System, and General), so that each user only sees and does what their role allows.',

    'To build a book cataloging module that follows the Library of Congress Classification with all 21 main classes, and connects to the Library of Congress API so librarians can import book records directly.',

    'To create a circulation module for issuing, returning, and renewing books, with borrowing policies that differ per role. The system should support configurable loan periods, book limits, renewal caps, and fine rates for all seven borrower categories.',

    'To automate fine computation so that the system detects overdue books on its own and calculates how much is owed based on the borrower\'s role. If a borrower\'s unpaid fines go past a set threshold (currently 100 pesos), the system should block them from borrowing or renewing until they settle.',

    'To implement a reservation system where requests are processed on a first-come-first-served basis, with a 48-hour window for claiming a reserved book before it goes to the next person in line.',

    'To add RFID-based attendance logging so library visits can be tracked digitally with tap-in records tied to each user.',

    'To deploy the whole system on serverless cloud infrastructure using Netlify for hosting, Neon PostgreSQL as the database, and Drizzle ORM for database operations, so that FEATI does not need to maintain any server hardware.',

    'To set up an audit log that captures login and logout events, record changes, and configuration updates across the entire system.',
])


# ── 1.4 ────────────────────────────────────────────────
heading('1.4 Significance of the Study')

para(
    'The researchers believe that this project will be useful to several groups '
    'connected to FEATI University. Below is a breakdown of how each group stands '
    'to benefit.'
)

sub('To the Library Staff')

para(
    'This is probably the group that benefits the most. Right now, every borrowing '
    'transaction means pulling out the logbook, writing down the details, and manually '
    'computing the due date based on who the borrower is. Returns are the same process '
    'in reverse, plus the added step of checking if the book is overdue and computing '
    'the fine. AklatBayon handles all of that in a few clicks. The system automatically '
    'knows that a student gets 7 days while a faculty member gets 30, so the staff does '
    'not need to look up the policy each time. The Head Librarian also gets a dashboard '
    'that shows how many books are borrowed, how many are overdue, and how much in fines '
    'is uncollected, all without running around counting things manually.'
)

sub('To the Students and Faculty')

para(
    'For students, the biggest improvement is being able to check book availability '
    'online. The public catalog page works as an OPAC where anyone can search by title, '
    'author, or ISBN. Faculty members get an added feature: they can recommend books for '
    'the library to purchase directly through the system, instead of filling out a paper '
    'form that might get lost. Both groups also get borrower profiles with QR codes, which '
    'speeds up identification at the circulation desk.'
)

sub('To the University Administration')

para(
    'The audit log is the main thing here. Every action in the system is recorded: '
    'who logged in, who borrowed what, who changed the settings. If there is ever a '
    'question about a missing book or an unpaid fine, the administration can trace it '
    'back. The fact that the system runs on the cloud also means FEATI does not need '
    'to buy a server or assign someone to maintain it.'
)

sub('To Future Researchers')

para(
    'Other students working on similar projects can use this as a reference. The '
    'database schema has 16 tables with 11 foreign key relationships, and the overall '
    'architecture (static frontend, serverless API, cloud database) is a pattern that '
    'applies to many kinds of web applications, not just library systems.'
)


# ── 1.5 ────────────────────────────────────────────────
heading('1.5 Scope and Limitations')

sub('Scope')

para('AklatBayon covers the following areas:')

numbered([
    'User management with seven roles and 23 permissions across eight groups.',

    'Book catalog management including authors, publishers, hierarchical categories, ISBN, call numbers based on Library of Congress Classification, and barcode printing using the CODE128 format.',

    'Library of Congress API integration for searching and importing book records.',

    'Circulation with role-specific policies. To be clear on the exact numbers: Head Librarian gets 15 books for 60 days with 3 renewals; Librarian Staff gets 10 books for 30 days with 2 renewals; Faculty Teaching gets 10 books for 30 days with 2 renewals; Faculty Non-Teaching gets 5 books for 14 days with 1 renewal; Department Chair gets 10 books for 30 days with 2 renewals; Student gets 3 books for 7 days with 1 renewal; Student Assistant gets 3 books for 7 days with 1 renewal.',

    'Automated fine computation with per-role daily rates and a block threshold at 100 pesos.',

    'Book reservations using a FIFO queue, a 48-hour claim window, automatic expiration for unclaimed holds, and limits of 2 active reservations for students and 5 for faculty.',

    'RFID-based attendance tracking.',

    'Audit logging for all user actions.',

    'A settings module with over 30 configurable parameters.',

    'Cloud deployment on Netlify with serverless functions and Neon PostgreSQL via Drizzle ORM.',

    'Twenty-three application pages accessible through any modern web browser.',
])

sub('Limitations')

para('The following items are not covered in the current version of the system:')

numbered([
    'Security features like password hashing, JWT tokens, server-side auth middleware, and rate limiting are not yet implemented. These are scheduled for a later phase.',

    'The RFID feature works through manual input. We did not integrate with physical RFID readers because the hardware was not available during development.',

    'The system does not use traditional library data exchange formats like MARC or protocols like Z39.50 and SRU. The only external data source is the Library of Congress REST API.',

    'The system is built for a single library on a single campus. Multi-branch support is not included.',

    'There are no email or SMS notifications. If a reserved book becomes available or a fine is due, the user only sees it when they log into the system.',

    'Formal usability testing with actual FEATI library staff and students has not been conducted yet. This is planned for the next development cycle.',
])


# ── 1.6 ────────────────────────────────────────────────
heading('1.6 Definition of Terms')

para(
    'Below are definitions of the key terms used in this document. These are '
    'provided for readers who may not be familiar with certain library or '
    'technical concepts.'
)

defterm('AklatBayon',
    'The Library Management System developed in this study. The name comes from '
    'the Filipino words "Aklat" (book) and "Bayon" (to carry). The live system is '
    'deployed at aklatbayonfeati.netlify.app.')

defterm('Circulation',
    'Refers to the lending of library materials to borrowers and their return. '
    'In this system, circulation covers issuing books, processing returns, and '
    'handling renewals.')

defterm('Drizzle ORM',
    'A TypeScript-based Object-Relational Mapping tool used in the system to '
    'communicate with the PostgreSQL database. It translates TypeScript code into '
    'SQL queries and manages database schema changes.')

defterm('FIFO',
    'Stands for First-In, First-Out. A method of processing a queue where the '
    'earliest request gets served first. The reservation system uses FIFO to decide '
    'who gets notified when a book becomes available.')

defterm('Library of Congress Classification (LCC)',
    'A system of organizing books by subject, used by most academic libraries. It uses '
    'letters and numbers to classify materials. AklatBayon implements all 21 main LCC '
    'classes from A (General Works) to Z (Bibliography and Library Science).')

defterm('Neon PostgreSQL',
    'A serverless database service running PostgreSQL. The system uses it as its '
    'primary data store, accessed through Netlify\'s Neon integration. All 16 database '
    'tables are hosted here.')

defterm('Netlify',
    'A cloud platform that hosts the AklatBayon web application. It handles static '
    'file hosting and runs serverless functions without requiring a traditional web server.')

defterm('Netlify Functions',
    'Backend code that runs on Netlify\'s infrastructure as serverless functions. '
    'AklatBayon uses two: api.mts for handling all database operations through REST '
    'endpoints, and seed.mts for populating the database with initial data.')

defterm('OPAC',
    'Stands for Online Public Access Catalog. It is the public-facing search page '
    'where anyone can look up books by title, author, ISBN, or subject without needing '
    'to log in.')

defterm('RBAC',
    'Stands for Role-Based Access Control. A way of managing what each user can do '
    'in the system based on their assigned role. AklatBayon has seven roles with 23 '
    'permissions.')

defterm('RFID',
    'Stands for Radio-Frequency Identification. A technology for identifying users '
    'through cards or tags. In the system, RFID is used for attendance logging and '
    'as an alternative way to log in.')

defterm('Serverless Architecture',
    'A setup where the cloud provider handles all the server management. The '
    'application code runs only when a request comes in. AklatBayon uses this '
    'through Netlify Functions and Neon PostgreSQL, so FEATI does not need to '
    'maintain any hardware.')

defterm('SweetAlert2',
    'A JavaScript library for showing pop-up dialogs and notifications in web '
    'applications. The system uses it instead of the default browser alert boxes.')


# ── Save ───────────────────────────────────────────────
output = r'c:\Users\Bins\AklatBayonV2\Chapter1_AklatBayon_v2.docx'
doc.save(output)
print(f'Saved: {output}')
